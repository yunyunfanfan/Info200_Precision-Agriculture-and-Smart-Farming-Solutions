from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches


ROOT = Path(__file__).resolve().parents[2]
TEMPLATE = ROOT / "Info 200 Team Ass Template.docx"
OUTPUT = ROOT / "Info 200 Team Ass Filled Draft.docx"
FIG_DIR = ROOT / "report_assets" / "figures"


TEAM_INFO_LINES = [
    "Team Number: XXX",
    "Team Leader: Name (English and Chinese) [ID]",
    "Team Members:",
    "Name 1 (English and Chinese) [ID], Name 2 (English and Chinese) [ID], Name 3 (English and Chinese) [ID]",
    "Emails: email1@example.com, email2@example.com, email3@example.com",
]


def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None


def clear_paragraph(paragraph):
    p = paragraph._element
    for child in list(p):
        if child.tag != qn("w:pPr"):
            p.remove(child)


def add_field(paragraph, instruction, display_text):
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction

    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")

    text = OxmlElement("w:t")
    text.text = display_text

    result_run = OxmlElement("w:r")
    result_run.append(text)

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")

    r1 = OxmlElement("w:r")
    r1.append(begin)
    r2 = OxmlElement("w:r")
    r2.append(instr)
    r3 = OxmlElement("w:r")
    r3.append(separate)
    r4 = OxmlElement("w:r")
    r4.append(end)

    paragraph._p.append(r1)
    paragraph._p.append(r2)
    paragraph._p.append(r3)
    paragraph._p.append(result_run)
    paragraph._p.append(r4)


def set_cover_page(doc):
    cell = doc.tables[0].cell(0, 0)
    for paragraph in cell.paragraphs[1:][::-1]:
        delete_paragraph(paragraph)
    paragraph = cell.paragraphs[0]
    clear_paragraph(paragraph)
    paragraph.style = doc.styles["No Spacing"]
    for i, line in enumerate(TEAM_INFO_LINES):
        run = paragraph.add_run(line)
        if i < len(TEAM_INFO_LINES) - 1:
            run.add_break()


def reset_body(doc):
    for paragraph in doc.paragraphs[2:][::-1]:
        delete_paragraph(paragraph)


def add_heading(doc, text):
    p = doc.add_paragraph(style="Heading 1")
    p.add_run(text)
    return p


def add_paragraph(doc, text="", style="Normal", align=None):
    p = doc.add_paragraph(style=style)
    if text:
        p.add_run(text)
    if align is not None:
        p.alignment = align
    return p


def add_bullets(doc, items):
    for item in items:
        add_paragraph(doc, item, style="List Paragraph")


def add_numbered_lines(doc, items):
    for idx, item in enumerate(items, start=1):
        add_paragraph(doc, f"{idx}. {item}", style="Normal")


def add_figure(doc, filename, caption):
    p = add_paragraph(doc, style="Normal", align=WD_ALIGN_PARAGRAPH.CENTER)
    run = p.add_run()
    run.add_picture(str(FIG_DIR / filename), width=Inches(6.2))
    caption_p = add_paragraph(doc, caption, style="Normal", align=WD_ALIGN_PARAGRAPH.CENTER)
    return p, caption_p


def add_simple_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for cell, value in zip(hdr, headers):
        cell.text = value
    for row in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row):
            cell.text = value
    return table


def build_document():
    doc = Document(str(TEMPLATE))
    set_cover_page(doc)

    doc.paragraphs[0].text = "Report Title: Designing a Data-Driven Precision Agriculture Information System for Regional Smart Farming"
    doc.paragraphs[1].text = "24/04/2026"

    reset_body(doc)

    add_heading(doc, "Contents")
    toc = add_paragraph(doc, style="Normal")
    add_field(toc, 'TOC \\\\o "1-3" \\\\h \\\\z \\\\u', "Right-click and update the field to generate the table of contents.")
    doc.add_page_break()

    add_heading(doc, "Executive Summary")
    add_paragraph(
        doc,
        "This report presents a formal proposal for the design of a precision agriculture information system for a regional agricultural cooperative. "
        "The current operating environment is characterised by fragmented field data, delayed decision-making, inefficient use of resources, and weak "
        "coordination between production and logistics activities. Existing farming practices remain heavily dependent on manual observation, historical "
        "routines, and loosely connected communication channels, thereby limiting the cooperative's ability to respond effectively to changing soil "
        "conditions, weather patterns, and crop-health risks.",
    )
    add_paragraph(
        doc,
        "The proposed solution is an integrated smart farming platform that combines IoT-based soil sensing, drone-assisted crop imaging, external weather "
        "and market data, and a unified dashboard for farmers, agronomists, and cooperative managers. The proposed system is intended to support real-time "
        "monitoring, irrigation planning, fertiliser optimisation, pest-risk alerts, and logistics coordination. In addition, it incorporates role-based "
        "access control, secure cloud storage, and auditable data governance mechanisms to address operational reliability, privacy, and ethical responsibility.",
    )
    add_paragraph(
        doc,
        "The report analyses the current organisational and technical context, identifies stakeholder requirements through interviews and questionnaires, "
        "and develops a set of practical recommendations supported by data flow diagrams, a data dictionary, process specifications, and a prototype dashboard "
        "concept. The analysis indicates that the proposed information system could improve water-use efficiency, reduce unnecessary chemical inputs, enhance "
        "visibility across the agricultural supply chain, and contribute to more sustainable, resilient, and economically viable farming operations.",
    )
    add_figure(doc, "figure_1_problem_chart.png", "Figure 1. Major weaknesses in the current farming model.")

    add_heading(doc, "Introduction")
    add_paragraph(
        doc,
        "Precision agriculture refers to the use of sensing technologies, data analytics, digital communication systems, and decision-support tools to manage "
        "agricultural inputs and operations with greater accuracy and responsiveness. Rather than relying on uniform treatment across all fields according to a "
        "fixed seasonal schedule, smart farming enables decisions to be adapted to local soil conditions, crop stress indicators, weather variability, and changing market demands.",
    )
    add_paragraph(
        doc,
        "For regional agricultural cooperatives, the transition from traditional management practices to smart farming is not merely a technological upgrade; "
        "it is also a significant information systems challenge. Cooperative members often operate across geographically dispersed plots, maintain inconsistent "
        "forms of record-keeping, and rely on limited integration between planting, irrigation, agronomy, harvesting, and transportation activities. "
        "Consequently, many operational decisions are delayed, reactive, or based on incomplete and outdated information.",
    )
    add_paragraph(
        doc,
        "This report therefore focuses on the design of an integrated information system capable of supporting a regional agricultural cooperative in its "
        "movement towards a more responsive, efficient, and sustainable operating model. The report is aligned with the principal themes of the Info 200 course, "
        "including systems analysis, organisational modelling, project planning, information gathering, process modelling, and structured decision-making. "
        "It aims not only to propose a technical solution, but also to demonstrate how information systems thinking can be applied to a complex real-world agricultural setting.",
    )

    add_heading(doc, "Analysis of the Problem")
    add_paragraph(doc, "Current Organisational Situation")
    add_paragraph(
        doc,
        "The present farming model can be described as semi-manual, fragmented, and only weakly integrated. Field observations are commonly recorded through "
        "paper notes, telephone communication, or informal spreadsheets rather than through a standardised digital system. Irrigation and fertiliser schedules "
        "are typically determined by previous experience or fixed timetables rather than by sensor-based evidence. Drone or satellite imaging, where available, "
        "is used only occasionally and is not incorporated into routine decision-making. Furthermore, logistics information is largely separated from crop-condition data, "
        "with the result that harvesting and delivery coordination are often reactive rather than strategically planned.",
    )
    add_paragraph(
        doc,
        "The cooperative structure creates additional organisational complexity. Individual farmers may differ substantially in terms of digital literacy, land quality, "
        "crop type, and management priorities. Agronomists and cooperative managers therefore require a system capable of standardising essential data while remaining "
        "sufficiently flexible to reflect local field realities and practical working conditions.",
    )
    add_paragraph(doc, "Strengths of the Existing Model")
    add_bullets(
        doc,
        [
            "Farmers possess valuable tacit knowledge concerning local climate conditions, soil behaviour, and crop cycles.",
            "The cooperative already has an existing communication network and recognised operational authority.",
            "Current working routines provide a useful baseline from which redesigned processes can emerge.",
            "The existing model is familiar to users and therefore requires little immediate technical training.",
        ],
    )
    add_paragraph(doc, "Limitations of the Existing Model")
    add_bullets(
        doc,
        [
            "Data collection is inconsistent, decentralised, and rarely available in real time.",
            "Irrigation and fertiliser decisions are not sufficiently evidence-based.",
            "Crop stress, disease outbreaks, and pest conditions may be detected too late for effective intervention.",
            "Managers are unable to combine field data, inventory information, and logistics schedules within a single decision environment.",
            "Independent farmers may have legitimate concerns regarding data ownership, surveillance, and secondary data use.",
        ],
    )
    add_paragraph(doc, "Stakeholder Analysis")
    add_bullets(
        doc,
        [
            "Farmers require simple mobile access, timely alerts, and recommendations that can be translated into immediate field action.",
            "Agronomists require accurate field data, image analysis capabilities, and comparison tools for diagnosis and intervention planning.",
            "Cooperative managers require dashboards, planning tools, and performance reporting to support oversight and resource allocation.",
            "Logistics providers require better forecasts of harvest volume, timing, and delivery windows.",
            "Regulators and consumers benefit indirectly from improved traceability, sustainability reporting, and more accountable data handling.",
        ],
    )
    add_paragraph(doc, "Information Gathering Design")
    add_paragraph(
        doc,
        "In order to understand stakeholder needs in a rigorous manner, the team should combine interactive and unobtrusive information-gathering methods. "
        "This mixed-methods approach is appropriate because it captures both expressed user expectations and observable operational realities.",
    )
    add_paragraph(doc, "Interview Questions")
    add_bullets(
        doc,
        [
            "Farmers: Which decisions are most difficult during irrigation, fertilisation, and pest control?",
            "Agronomists: Which data is currently missing when diagnosing crop problems?",
            "Managers: Which reports are most necessary for planning, supervision, and performance control?",
            "Logistics providers: What information gaps most commonly delay transport preparation?",
        ],
    )
    add_paragraph(doc, "Questionnaire Items")
    add_numbered_lines(
        doc,
        [
            "How often do you use digital tools in farm operations?",
            "Which problems cause the greatest production loss?",
            "How useful would real-time moisture alerts be?",
            "Which device do you prefer for system access: mobile phone, tablet, or desktop?",
            "How concerned are you about privacy and data ownership?",
        ],
    )
    add_paragraph(doc, "Unobtrusive Methods")
    add_bullets(
        doc,
        [
            "Review of existing field logs, irrigation records, and agronomic notes",
            "Observation of seasonal workflows and routine coordination practices",
            "Analysis of communication delays between farms and the cooperative office",
            "Inspection of current spreadsheets, paper forms, and reporting templates",
        ],
    )

    add_heading(doc, "Proposals and Solutions")
    add_paragraph(
        doc,
        "The proposed solution is a unified smart farming information system that gathers data from field sensors, drone imagery, and external data services, "
        "and then transforms that data into operationally meaningful recommendations. The proposed architecture is intentionally modular so that the cooperative "
        "can begin with a limited set of high-value functions, such as irrigation monitoring, before gradually expanding the system to include fertiliser management, "
        "pest detection, traceability, and market planning.",
    )
    add_figure(doc, "figure_2_architecture.png", "Figure 2. Proposed smart farming system architecture.")
    add_paragraph(doc, "Core Functional Components")
    add_bullets(
        doc,
        [
            "IoT sensing layer: soil moisture, soil temperature, pH, humidity, and weather sensors installed across representative field zones.",
            "Aerial monitoring layer: periodic drone imagery to identify crop stress, disease patterns, and irrigation gaps.",
            "Data integration layer: an IoT gateway and cloud platform for validating, storing, and combining data streams.",
            "Decision-support layer: rules and analytics for irrigation scheduling, fertiliser planning, anomaly detection, and yield forecasting.",
            "Presentation layer: mobile and web dashboards tailored to farmers, agronomists, and managers.",
        ],
    )
    add_paragraph(doc, "Recommended Innovations")
    add_numbered_lines(
        doc,
        [
            "A field-zoning approach so that each plot receives recommendations based on actual conditions rather than a uniform schedule.",
            "Automated threshold alerts for low soil moisture, disease risk, and abnormal sensor readings.",
            "A prioritised dashboard that presents urgent operational actions rather than raw data alone.",
            "Shared visibility between production and logistics teams so that harvesting and transport can be coordinated earlier.",
            "Traceable digital records to support sustainability reporting, accountability, and continuous improvement.",
        ],
    )
    add_paragraph(doc, "User-Centred Design Features")
    add_bullets(
        doc,
        [
            "Colour-coded alerts",
            "Recommended actions",
            "Field-by-field status",
            "Offline-capable mobile access",
            "Large buttons and low-text interaction patterns",
            "Bilingual interface support where appropriate",
        ],
    )

    add_heading(doc, "Evaluation of Solutions")
    add_paragraph(doc, "Feasibility")
    add_paragraph(
        doc,
        "The proposal is feasible because it can be implemented incrementally. The cooperative does not need to digitise every function simultaneously. "
        "A pilot phase can begin with a subset of farms and a limited set of use cases, such as irrigation optimisation and field-health monitoring.",
    )
    add_paragraph(
        doc,
        "From a technical perspective, the required technologies are already available and sufficiently mature for practical deployment. Sensor hardware, cloud services, "
        "web dashboards, and API-based weather integration are widely accessible. The principal challenge therefore lies not in technical possibility, but in user adoption, "
        "staff training, governance, and alignment with existing operational routines.",
    )
    add_paragraph(doc, "Expected Benefits")
    add_bullets(
        doc,
        [
            "Improved water-use efficiency through condition-based irrigation",
            "Better fertiliser targeting and lower environmental impact",
            "Faster response to pests, disease, and field anomalies",
            "Improved harvest planning and logistics coordination",
            "Stronger evidential support for management decisions and future investment",
        ],
    )
    add_paragraph(doc, "Risks and Mitigation")
    add_simple_table(
        doc,
        ["Risk", "Impact", "Mitigation"],
        [
            ["Sensor failure or poor calibration", "Inaccurate recommendations", "Scheduled maintenance and validation checks"],
            ["Low user adoption", "Limited business value", "Training, simple interfaces, pilot testing, and feedback loops"],
            ["Weak internet connectivity", "Data delay in rural areas", "Edge buffering and periodic synchronisation"],
            ["Privacy concerns", "Resistance from farmers", "Clear data governance and role-based access controls"],
            ["Cost pressure", "Slow rollout", "Phased implementation and cooperative-level investment"],
        ],
    )
    add_paragraph(doc, "Ethical Considerations")
    add_paragraph(
        doc,
        "Ethical considerations are central to the design of the proposed system. Farmers should understand what data is collected, why it is collected, who can access it, "
        "and how it may be used. Data should not be exploited in ways that weaken farmer autonomy or confer unfair advantage upon particular stakeholders. In addition, the "
        "system should avoid unnecessary algorithmic opacity by presenting recommendations in clear language and, wherever possible, by showing the indicators or thresholds "
        "on which those recommendations are based.",
    )

    add_heading(doc, "Design and Technology")
    add_paragraph(doc, "Web 2.0 and Web 3.0 Perspective")
    add_paragraph(
        doc,
        "The proposed platform is grounded primarily in Web 2.0 principles, including cloud-hosted services, user accounts, dashboards, interactive analytics, and collaborative "
        "information sharing. Nevertheless, selected Web 3.0 concepts may also be relevant, particularly in relation to traceability and trusted data exchange. For example, the "
        "cooperative could in future investigate tamper-evident records for crop origin, treatment history, or certification purposes if regulatory pressures or market requirements "
        "justify the additional complexity.",
    )
    add_paragraph(doc, "Prototype Dashboard Concept")
    add_paragraph(
        doc,
        "The dashboard should combine monitoring, analysis, and action within a single interface. Rather than requiring users to navigate multiple disconnected tools, the design "
        "should present key indicators, urgent alerts, map-based status information, and direct links to external services such as weather APIs and market-price resources. "
        "This approach reflects the principle that an effective information system should reduce cognitive burden while improving operational visibility.",
    )
    add_figure(doc, "figure_4_dashboard.png", "Figure 3. Dashboard wireframe for farmers and cooperative managers.")
    add_paragraph(doc, "Data Flow Diagram")
    add_paragraph(
        doc,
        "The information system must support the structured movement of data from farm actors and automated sensing tools into processing functions, data stores, and shared decision outputs. "
        "The following Level-1 DFD illustrates the principal data flows among users, system processes, repositories, and the final decision dashboard.",
    )
    add_figure(doc, "figure_3_dfd.png", "Figure 4. Level-1 data flow diagram for the proposed information system.")
    add_paragraph(doc, "Data Dictionary")
    add_simple_table(
        doc,
        ["Data Element", "Description", "Example", "Source"],
        [
            ["field_id", "Unique identifier for each field zone", "F-12", "Farm records"],
            ["soil_moisture", "Percentage moisture level in soil", "24%", "IoT sensor"],
            ["crop_health_index", "Calculated crop condition score", "0.81", "Drone analytics"],
            ["irrigation_status", "Current irrigation state", "Active / Scheduled", "Control system"],
            ["alert_level", "Priority of warning or issue", "High", "Analytics engine"],
            ["delivery_window", "Planned transport period", "15:00-18:00", "Logistics provider"],
        ],
    )
    add_paragraph(doc, "Process Specification Example")
    add_paragraph(doc, "Process 2.0: Analyse Conditions")
    add_bullets(
        doc,
        [
            "Input: sensor readings, drone imagery, weather forecast data, and field records",
            "Validate incoming values and remove invalid records.",
            "Compare current values with crop thresholds and historical patterns.",
            "Detect abnormal conditions such as water stress or sudden temperature change.",
            "Produce recommended actions and alert levels.",
            "Output: recommendations, alerts, and dashboard summaries",
        ],
    )

    add_heading(doc, "Methodology")
    add_paragraph(doc, "Development Approach")
    add_paragraph(
        doc,
        "The project should follow an iterative systems analysis and design approach. Early activities should concentrate on understanding the organisational context, defining the "
        "principal decision problems, and identifying the minimum viable requirements for a pilot deployment. A prototype should then be tested with a small group of representative "
        "users before any wider rollout is attempted.",
    )
    add_paragraph(doc, "Research Methods")
    add_bullets(
        doc,
        [
            "Qualitative interviews to identify stakeholder expectations and operational pain points",
            "Questionnaires to compare requirements across user groups",
            "Document analysis to review existing forms, records, and workflows",
            "Structured modelling to describe processes and information flows",
        ],
    )
    add_paragraph(doc, "Suggested Timescale")
    add_figure(doc, "figure_5_gantt.png", "Figure 5. Suggested project timescale for the team assignment.")
    add_paragraph(doc, "Security, Storage, and Access Control")
    add_bullets(
        doc,
        [
            "Cloud-based storage for scalability and controlled shared access",
            "Encrypted data transfer between sensors, gateways, and servers",
            "Role-based access control for farmers, agronomists, managers, and administrators",
            "Audit logs for sensitive data access and administrative actions",
            "Retention rules aligned with legal and regulatory requirements",
            "Backup and recovery procedures to support operational continuity",
        ],
    )

    add_heading(doc, "Conclusion and Recommendations")
    add_paragraph(
        doc,
        "The analysis undertaken in this report demonstrates that the cooperative's current farming model is constrained by fragmented information, manual coordination practices, "
        "and a limited capacity for timely data-driven action. A precision agriculture information system offers a credible pathway towards more efficient and sustainable farm "
        "management by integrating sensing, analytics, and decision support within a coherent operational platform.",
    )
    add_paragraph(
        doc,
        "The principal recommendation is that implementation should proceed in phases. The first phase should focus on a pilot deployment involving selected fields, moisture sensors, "
        "weather-data integration, and a simplified dashboard for irrigation planning and alert management. Once practical value has been demonstrated and user feedback has been "
        "incorporated, the cooperative can extend the platform to include drone analytics, fertiliser optimisation, supply-chain visibility, and more advanced forecasting functions.",
    )
    add_paragraph(
        doc,
        "In conclusion, smart farming should not be pursued as technology for its own sake. It should instead be designed as a user-centred information system that improves everyday "
        "agricultural decision-making, strengthens environmental sustainability, and supports the long-term resilience of the farming community.",
    )

    doc.save(str(OUTPUT))
    print(f"Generated: {OUTPUT}")


if __name__ == "__main__":
    build_document()
